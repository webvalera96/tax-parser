import io
import csv
from docx import Document
import pandas as pd


def read_docx_tables(filename, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)
    Parameters:
        
        filename:   file name of a Word Document
        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)        
        kwargs:     arguments to pass to `pd.read_csv()` function
    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """

    def read_docx_tab(tab, **kwargs):        
        vf = io.StringIO()
        writer = csv.writer(vf)       
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)   
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise